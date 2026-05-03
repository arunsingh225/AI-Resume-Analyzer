import io
import re
import datetime
from typing import Tuple, Dict

import pdfplumber
import docx

from app.constants import ACTION_VERBS, determine_level


# ─── Text Cleaning ──────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    import unicodedata
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'[^\w\s\.\+\#\&/\-\•\*\–\—:,;@()]', ' ', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Insert newline before ALL-CAPS headings that pdfplumber glues to content
    # e.g., "...some textPROFESSIONAL SUMMARY" → "...some text\nPROFESSIONAL SUMMARY"
    text = re.sub(
        r'(?<=[a-z0-9\.\)])([A-Z]{2,}(?:\s+[A-Z]+)*(?:\s*(?:\&|AND)\s*[A-Z]+)*)\s*\n',
        r'\n\1\n',
        text
    )
    return text.strip()


def extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:6]:
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
            if not any(ch.isdigit() for ch in line) and '@' not in line:
                return line
    return "Candidate"


def extract_email(text: str) -> str:
    m = re.search(r'[\w\.\+\-]+@[\w\.\-]+\.\w{2,}', text)
    return m.group(0) if m else ""


def extract_phone(text: str) -> str:
    m = re.search(r'(\+?\d[\d\s\-\(\)]{7,16}\d)', text)
    return m.group(0).strip() if m else ""


def extract_years_experience(text: str) -> float:
    text = text.lower()

    # education section ke baad ka text ignore karo
    text = re.split(r'\beducation\b', text)[0]

    # direct patterns first
    direct_patterns = [
        r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years|year|yrs|yr)\s+of\s+(?:total\s+)?(?:work\s+)?experience',
        r'experience\s+of\s+(\d+(?:\.\d+)?)\s*\+?\s*(?:years|year|yrs|yr)',
    ]
    for pat in direct_patterns:
        m = re.search(pat, text)
        if m:
            return float(m.group(1))

    # months
    month_match = re.search(r'(\d+(?:\.\d+)?)\s*(months|month|mos|mo)\b', text)
    if month_match:
        months = float(month_match.group(1))
        return round(months / 12, 1)

    # sum years + months only from experience section
    exp_section = text
    match = re.search(r'(experience|work experience)(.*?)(education|projects|skills|certifications|$)', text, re.DOTALL)
    if match:
        exp_section = match.group(2)

    years = re.findall(r'(\d+(?:\.\d+)?)\s*(?:years|year|yrs|yr)\b', exp_section)
    months = re.findall(r'(\d+(?:\.\d+)?)\s*(?:months|month|mos|mo)\b', exp_section)

    total = sum(float(y) for y in years) + sum(float(m) / 12 for m in months)
    if total > 0:
        return round(total, 1)

    # date range fallback
    year_ranges = re.findall(r'(20\d{2})\s*[-–—]\s*(20\d{2}|present|current|now)', exp_section)
    if year_ranges:
        now = datetime.datetime.now().year
        total = 0
        for start, end in year_ranges:
            s = int(start)
            e = now if end in ('present', 'current', 'now') else int(end)
            total += max(0, e - s)
        return min(total, 30)

    return 0.0
    





# ─── PDF Extraction ──────────────────────────────────────────────────────────

def extract_from_pdf(file_bytes: bytes) -> Tuple[str, Dict]:
    pages = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
          for page in pdf.pages:
            t = page.extract_text(x_tolerance=2, y_tolerance=2)

            # 👇 IMPORTANT FIX
            if not t:
                t = page.extract_text(layout=True)

            if t:
                pages.append(t)
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {e}")

    if not pages:
        raise ValueError(
            "No extractable text found. The PDF may be image-only or scanned. "
            "Please use a text-based PDF resume."
        )

    raw = "\n".join(pages)
    text = clean_text(raw)
    years = extract_years_experience(text)
    meta = {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "years_experience": years,
        "experience_level": determine_level(years),
        "file_type": "PDF",
    }
    return text, meta


# ─── DOCX Extraction ─────────────────────────────────────────────────────────

def extract_from_docx(file_bytes: bytes) -> Tuple[str, Dict]:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")

    paragraphs = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            paragraphs.append(t)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    paragraphs.append(t)

    if not paragraphs:
        raise ValueError("No text found in the DOCX file.")

    raw = "\n".join(paragraphs)
    text = clean_text(raw)
    years = extract_years_experience(text)
    meta = {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "years_experience": years,
        "experience_level": determine_level(years),
        "file_type": "DOCX",
    }
    return text, meta


# ─── Formatting Quality ───────────────────────────────────────────────────────




def check_formatting(text: str) -> Dict:
    tl = text.lower()
    score = 0
    checks = {}

    checks["has_email"] = bool(re.search(r'[\w\.\+]+@[\w\.\-]+\.\w+', text))
    checks["has_phone"] = bool(re.search(r'(\+?\d[\d\s\-]{7,16}\d)', text))
    checks["has_linkedin"] = bool(re.search(r'linkedin\.com', tl))
    checks["has_github_or_portfolio"] = bool(re.search(r'github\.com|gitlab\.com|dribbble\.com|behance\.net|portfolio', tl))
    checks["has_professional_links"] = checks["has_linkedin"] or checks["has_github_or_portfolio"]
    checks["has_bullets"] = bool(re.search(r'[•\-\*]\s+\w', text))
    checks["has_action_verbs"] = sum(1 for v in ACTION_VERBS if v in tl) >= 4
    checks["has_dates"] = bool(re.search(r'\b(20\d{2}|19\d{2})\b', text))
    checks["has_quantified"] = bool(re.search(r'\d+\s*(%|percent|million|billion|lakh|crore|\bk\b|users|clients|projects)', tl))
    checks["clean_chars"] = len(re.findall(r'[^\x00-\x7F]', text)) < 30
    checks["reasonable_length"] = 200 < len(text.split()) < 1200

    weights = {
        "has_email": 15, "has_phone": 10,
        "has_professional_links": 12,
        "has_linkedin": 3,
        "has_github_or_portfolio": 3,
        "has_bullets": 18,
        "has_action_verbs": 18, "has_dates": 10, "has_quantified": 10,
        "clean_chars": 3, "reasonable_length": 3,
    }
    score = sum(weights[k] for k, v in checks.items() if v)
    return {"score": min(score, 100), "checks": checks}
